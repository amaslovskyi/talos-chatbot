"""
Document loader and chunking module.
Handles loading various document types and splitting them into chunks for vector storage.
"""

import os
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path

# Document processing imports
import pypdf
from docx import Document
from bs4 import BeautifulSoup
import markdown

# LangChain imports for text splitting
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument

from config import get_settings

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentLoader:
    """
    Loads and processes documents from a directory.
    Supports PDF, DOCX, TXT, MD, and HTML files.
    """

    def __init__(self):
        """Initialize the document loader with settings."""
        self.settings = get_settings()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.settings.chunk_size,
            chunk_overlap=self.settings.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""],
        )

    def load_all_knowledge_bases(self) -> List[LangChainDocument]:
        """
        Load documents from all configured knowledge base paths.

        Returns:
            List of LangChain Document objects from all knowledge bases.
        """
        all_documents = []

        # Load from primary documents directory
        primary_docs = self.load_documents(self.settings.documents_directory)
        all_documents.extend(primary_docs)
        logger.info(
            f"Loaded {len(primary_docs)} documents from primary directory: {self.settings.documents_directory}"
        )

        # Load from additional knowledge base paths
        if self.settings.additional_knowledge_paths:
            additional_paths = [
                path.strip()
                for path in self.settings.additional_knowledge_paths.split(",")
                if path.strip()
            ]

            for path in additional_paths:
                if os.path.exists(path):
                    kb_docs = self.load_documents(path)
                    all_documents.extend(kb_docs)
                    logger.info(
                        f"Loaded {len(kb_docs)} documents from knowledge base: {path}"
                    )
                else:
                    logger.warning(f"Knowledge base path does not exist: {path}")

        logger.info(
            f"Total documents loaded from all knowledge bases: {len(all_documents)}"
        )
        return all_documents

    def load_documents(
        self, directory: Optional[str] = None
    ) -> List[LangChainDocument]:
        """
        Load all supported documents from the specified directory.

        Args:
            directory: Path to documents directory. Uses config default if None.

        Returns:
            List of LangChain Document objects with content and metadata.
        """
        if directory is None:
            directory = self.settings.docs_directory

        documents = []
        doc_path = Path(directory)

        if not doc_path.exists():
            logger.warning(f"Documents directory does not exist: {directory}")
            return documents

        # Supported file extensions
        supported_extensions = {".pdf", ".docx", ".txt", ".md", ".html", ".htm"}

        # Walk through directory and process files
        for file_path in doc_path.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                try:
                    logger.info(f"Processing file: {file_path}")
                    doc_content = self._load_single_document(file_path)

                    if doc_content:
                        # Create LangChain document with metadata
                        doc = LangChainDocument(
                            page_content=doc_content,
                            metadata={
                                "source": str(file_path),
                                "filename": file_path.name,
                                "file_type": file_path.suffix.lower(),
                                "file_size": file_path.stat().st_size,
                            },
                        )
                        documents.append(doc)

                except Exception as e:
                    logger.error(f"Error processing file {file_path}: {str(e)}")
                    continue

        logger.info(f"Loaded {len(documents)} documents")
        return documents

    def _load_single_document(self, file_path: Path) -> Optional[str]:
        """
        Load content from a single document file.

        Args:
            file_path: Path to the document file.

        Returns:
            Document content as string, or None if loading failed.
        """
        file_extension = file_path.suffix.lower()

        try:
            if file_extension == ".pdf":
                return self._load_pdf(file_path)
            elif file_extension == ".docx":
                return self._load_docx(file_path)
            elif file_extension == ".txt":
                return self._load_text(file_path)
            elif file_extension == ".md":
                return self._load_markdown(file_path)
            elif file_extension in [".html", ".htm"]:
                return self._load_html(file_path)
            else:
                logger.warning(f"Unsupported file type: {file_extension}")
                return None

        except Exception as e:
            logger.error(f"Error loading {file_path}: {str(e)}")
            return None

    def _load_pdf(self, file_path: Path) -> str:
        """Load content from PDF file."""
        text = ""
        with open(file_path, "rb") as file:
            pdf_reader = pypdf.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()

    def _load_docx(self, file_path: Path) -> str:
        """Load content from DOCX file."""
        doc = Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return "\n".join(text).strip()

    def _load_text(self, file_path: Path) -> str:
        """Load content from text file."""
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read().strip()

    def _load_markdown(self, file_path: Path) -> str:
        """Load content from Markdown file and convert to plain text."""
        with open(file_path, "r", encoding="utf-8") as file:
            md_content = file.read()
        # Convert markdown to HTML, then extract text
        html = markdown.markdown(md_content)
        soup = BeautifulSoup(html, "html.parser")
        return soup.get_text().strip()

    def _load_html(self, file_path: Path) -> str:
        """Load content from HTML file."""
        with open(file_path, "r", encoding="utf-8") as file:
            html_content = file.read()
        soup = BeautifulSoup(html_content, "html.parser")
        return soup.get_text().strip()

    def chunk_documents(
        self, documents: List[LangChainDocument]
    ) -> List[LangChainDocument]:
        """
        Split documents into smaller chunks for better retrieval.

        Args:
            documents: List of LangChain documents to chunk.

        Returns:
            List of chunked documents with preserved metadata.
        """
        chunked_docs = []

        for doc in documents:
            # Split the document into chunks
            chunks = self.text_splitter.split_text(doc.page_content)

            # Create new documents for each chunk
            for i, chunk in enumerate(chunks):
                # Add chunk index to metadata
                chunk_metadata = doc.metadata.copy()
                chunk_metadata["chunk_index"] = i
                chunk_metadata["total_chunks"] = len(chunks)

                chunked_doc = LangChainDocument(
                    page_content=chunk, metadata=chunk_metadata
                )
                chunked_docs.append(chunked_doc)

        logger.info(
            f"Created {len(chunked_docs)} chunks from {len(documents)} documents"
        )
        return chunked_docs


def load_and_chunk_documents(
    directory: Optional[str] = None,
) -> List[LangChainDocument]:
    """
    Convenience function to load and chunk documents in one step.

    Args:
        directory: Path to documents directory.

    Returns:
        List of chunked LangChain documents.
    """
    loader = DocumentLoader()
    documents = loader.load_documents(directory)
    return loader.chunk_documents(documents)


def load_and_chunk_specific_files(
    file_paths: List[str],
) -> List[LangChainDocument]:
    """
    Load and chunk specific files by their paths.

    Args:
        file_paths: List of file paths to process.

    Returns:
        List of chunked LangChain documents.
    """
    loader = DocumentLoader()
    documents = []

    for file_path in file_paths:
        try:
            path = Path(file_path)
            if path.exists() and path.is_file():
                doc_content = loader._load_single_document(path)
                if doc_content:
                    doc = LangChainDocument(
                        page_content=doc_content,
                        metadata={
                            "source": str(path),
                            "filename": path.name,
                            "file_type": path.suffix.lower(),
                            "file_size": path.stat().st_size,
                        },
                    )
                    documents.append(doc)
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            continue

    return loader.chunk_documents(documents)
